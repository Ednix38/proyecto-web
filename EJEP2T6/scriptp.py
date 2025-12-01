import random
from pathlib import Path
from docx import Document
from docx.shared import Inches
from playwright.sync_api import sync_playwright

# --- CONFIGURACIÓN ---
CARPETA_HTML = Path(r"C:\Users\Edwin\OneDrive\Escritorio\Universidad\Cuatrimestre No.4\Apps web\PaginaPersonal\EJEP2T6")  # Cambia a tu carpeta de HTML
WORD_SALIDA = "resultados.docx"

# Crear documento Word
doc = Document()
doc.add_heading("Resultados de Ejercicios HTML", level=0)

def procesar_html(file_path, browser):
    nombre = file_path.stem
    doc.add_heading(nombre, level=1)

    # --- 1. Código HTML ---
    with open(file_path, "r", encoding="utf-8") as f:
        html_code = f.read()
    doc.add_heading("Código HTML", level=2)
    doc.add_paragraph(html_code)

    # --- Abrir página nueva ---
    page = browser.new_page()
    resultado_img = f"{nombre}_resultado.png"
    prompt_state = {"detected": False}  # Diccionario para manejar estado

    # Handler para prompts
    def handle_dialog(dialog):
        prompt_state["detected"] = True
        valor_random = str(random.randint(0, 100))
        try:
            dialog.accept(valor_random)
            print(f"{nombre}: prompt aceptado con valor {valor_random}")
        except Exception as e:
            print(f"{nombre}: prompt ya manejado, ignorando. {e}")

    # Escuchar un diálogo a la vez
    page.on("dialog", handle_dialog)
    page.goto(f"file:///{file_path.resolve()}")
    
    # Screenshot del resultado final
    page.screenshot(path=resultado_img)
    doc.add_heading("Resultado Final", level=2)
    doc.add_picture(resultado_img, width=Inches(6))

    # Captura del prompt si hubo
    if prompt_state["detected"]:
        doc.add_heading("Prompt/Inputs", level=2)
        doc.add_picture(resultado_img, width=Inches(6))

    page.close()

# --- Ejecutar Playwright ---
with sync_playwright() as p:
    browser = p.chromium.launch()
    for html_file in CARPETA_HTML.glob("*.html"):
        procesar_html(html_file, browser)
    browser.close()

# Guardar Word
doc.save(WORD_SALIDA)
print(f"Word generado: {WORD_SALIDA}")
